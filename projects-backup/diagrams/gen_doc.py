"""
生成 AED 系统架构设计介绍文档 (Word .docx)
嵌入全部 5 张模型图 JPEG，含 MagicGrid 方法论讲解和图间追溯关系。
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from pathlib import Path

INPUT_DIR = Path(r"E:\.Claude Code Project\diagrams\AED_JPEG")
OUTPUT_PATH = Path(r"E:\.Claude Code Project\diagrams\AED_系统架构设计文档.docx")

# Diagram definitions: (filename, title_en, title_cn, section)
DIAGRAMS = [
    ("AED_UseCase.jpg",       "Use Case Diagram",        "用例图",    "Requirements"),
    ("AED_Activity.jpg",      "Activity Diagram",        "活动图",    "Behavior"),
    ("AED_BDD.jpg",           "Block Definition Diagram", "BDD 块定义图", "Structure"),
    ("AED_Parametric.jpg",    "Parametric Diagram",      "参数图",    "Parameters"),
    ("AED_Sequence.jpg",      "Sequence Diagram",        "时序图",    "Interaction"),
]

# ── Document Setup ──────────────────────────────────────────────────
doc = Document()

# Page setup — A4 landscape for better diagram fit
for section in doc.sections:
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

# ── Helper functions ─────────────────────────────────────────────────
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return h

def add_para(text, bold=False, size=10.5, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    if alignment is not None:
        p.alignment = alignment
    return p

def add_rich_para(segments):
    """segments: list of (text, bold, size)"""
    p = doc.add_paragraph()
    for text, bold, size in segments:
        run = p.add_run(text)
        run.font.name = "微软雅黑"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(size)
        run.bold = bold
    return p

def add_image(img_path, width_inches=8.0):
    """Insert a JPEG inline, centered."""
    if not Path(img_path).exists():
        add_para(f"[图片未找到: {img_path}]", size=9)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))

def add_table(headers, rows, col_widths=None):
    """Insert a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph()  # spacer
    return table


# ═══════════════════════════════════════════════════════════════════════
#  COVER
# ═══════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_para("AED 自动外灯系统", bold=True, size=26, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para("架构设计文档", bold=True, size=20, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("基于 MagicGrid 方法论 · SysML v2 建模", size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para("2026-05-28", size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("本文档以 AED（Automatic Exterior Lighting，自动外灯）系统为例，"
         "完整演示基于 MagicGrid 方法论的 SysML 系统建模过程。"
         "包含用例图、活动图、BDD 块定义图、参数图、时序图共 5 张模型图，"
         "覆盖 Requirements / Behavior / Structure / Parameters 四大支柱，"
         "并建立从需求到交互的完整追溯链。", size=10.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
#  1. MAGICGRID 方法论框架
# ═══════════════════════════════════════════════════════════════════════
add_heading_styled("1. MagicGrid 方法论框架", level=1)

add_para("MagicGrid 是 NoMagic（现 Dassault）发布的 SysML 建模方法手册，"
         "将系统工程活动组织为「域 × 支柱」矩阵。"
         "本文档的 5 张图覆盖了问题域（Problem Domain）到解决域（Solution Domain）的关键格子。", size=10.5)

add_table(
    ["域 (Domain)", "Req (需求)", "Beh (行为)", "Str (结构)", "Par (参数)"],
    [
        ["Problem (Black-box)", "B1: Stakeholder Needs", "B2: Use Cases  ← 用例图", "B3: System Context", "B4: MoE"],
        ["Problem (White-box)", "W1 (合并至B1)", "W2: Functional Analysis  ← 活动图", "W3: Logical Subsystems", "W4: Subsystem MoEs"],
        ["Solution", "S1: System Reqs", "S2: System Behavior", "S3: System Structure  ← BDD", "S4: System Params  ← 参数图"],
        ["Verification", "—", "—", "—", "← 时序图 交叉验证"],
    ],
)

add_para("五张图的定位关系：", bold=True, size=10.5)
add_para(
    "  用例图 (Requirements) — 定义系统功能边界和外部参与者，回答「系统要做什么？」\n"
    "  活动图 (Behavior) — 展开核心用例的控制流和数据流，回答「如何做？」\n"
    "  BDD (Structure) — 定义系统组件、属性和组合/引用关系，回答「谁来做？」\n"
    "  参数图 (Parameters) — 从活动图决策节点提取约束方程，回答「规则是什么？」\n"
    "  时序图 (Interaction) — 选取一个具体运行场景，实例化所有视角并验证一致性",
    size=10.0,
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
#  2. 用例图 — Requirements Pillar
# ═══════════════════════════════════════════════════════════════════════
add_heading_styled("2. 用例图 — 系统功能边界", level=1)
add_para("MagicGrid 定位: B2 · Problem Domain · Behavior 列", bold=True, size=9.5)
add_para("用例图定义了 AED 系统与外部参与者（Driver、Light Sensor、Rain Sensor、BCM、Instrument Cluster）之间的功能契约。"
         "共识别 10 个用例，覆盖自动灯光控制、手动超控、状态上报和故障诊断四大类功能。", size=10.5)

add_para("关键用例清单：", bold=True, size=10.5)

add_table(
    ["用例编号", "用例名称", "类型", "关联关系"],
    [
        ["UC01", "自动控制近光灯", "核心功能", "include UC03, UC04；被 UC08 extend"],
        ["UC02", "自动控制远光灯", "核心功能", "extend UC01（仅在近光开启时可用）"],
        ["UC03", "自动控制日行灯 (DRL)", "法规功能", "被 UC01 include"],
        ["UC04", "自动控制位置灯", "法规功能", "被 UC01 include"],
        ["UC05", "自动控制雾灯", "辅助功能", "extend UC01（需要雨量传感器参与）"],
        ["UC06", "转弯辅助照明", "辅助功能", "include UC01"],
        ["UC07", "回家/离家照明", "舒适功能", "extend UC01"],
        ["UC08", "手动超控", "安全功能", "extend UC01, UC02"],
        ["UC09", "灯光状态上报", "诊断功能", "关联 BCM、Instrument Cluster"],
        ["UC10", "故障诊断与降级", "安全功能", "include UC09"],
    ],
)

add_image(INPUT_DIR / "AED_UseCase.jpg", width_inches=9.0)
add_para("图 2-1   AED 系统用例图 (Use Case Diagram)", size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_para("设计要点：", bold=True, size=10.5)
add_para(
    "1. UC08（手动超控）独立为 extend 用例而非放在 UC01 内部——这是 ISO 26262 ASIL-A 的要求："
    "驾驶员必须能随时退出自动模式，且该操作不受软件故障影响。\n"
    "2. UC10（故障诊断与降级）覆盖所有用例——任一灯具的断路/短路/过温都会触发 DTC 并进入安全状态（至少保留位置灯）。\n"
    "3. Light Sensor 和 Rain Sensor 作为 <<system>> actor，表示它们是 AED 系统外部的独立 ECU（通过 LIN 总线通信），而非 AED 内部组件。",
    size=10.0,
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
#  3. 活动图 — Behavior Pillar
# ═══════════════════════════════════════════════════════════════════════
add_heading_styled("3. 活动图 — 自动近光灯控制流程", level=1)
add_para("MagicGrid 定位: W2 · Problem Domain (White-box) · Behavior 列", bold=True, size=9.5)
add_para("活动图选取用例图中最核心的 UC01（自动控制近光灯），展开为完整的控制流和数据流。"
         "泳道划分直接对应 BDD 中的物理 Block，数据流对应 Block 的属性。", size=10.5)

add_para("从用例图到活动图的映射：", bold=True, size=10.5)
add_table(
    ["用例图元素", "→", "活动图对应"],
    [
        ["UC01 自动控制近光灯", "→", "活动图主流程（自检→光照采集→决策→执行）"],
        ["Light Sensor actor", "→", "Swimlane: Light Sensor（光照采集 + 50ms 滤波）"],
        ["Rain Sensor actor", "→", "Swimlane: Rain Sensor（雨量采集）"],
        ["Driver actor", "→", "Swimlane: Driver（手动超控节点）"],
        ["UC08 手动超控 (extend UC01)", "→", "Driver swimlane 中超控仲裁"],
        ["UC10 故障诊断与降级", "→", "尾部故障检测分支（电流异常→DTC→替代策略）"],
    ],
)

add_image(INPUT_DIR / "AED_Activity.jpg", width_inches=9.0)
add_para("图 3-1   AED 自动近光灯控制活动图 (Activity Diagram)", size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_para("流程关键路径：", bold=True, size=10.5)
add_para(
    "Path A (正常夜间): 自检通过 → 光照 < 400 lux → 近光灯 ON → 状态上报 → 循环\n"
    "Path B (黄昏过渡): 自检通过 → 光照 400-1000 lux → DRL ON → 状态上报 → 循环\n"
    "Path C (雨天强制): 自检通过 → 雨量 >= 2 → 强制近光灯 ON（无论光照值）\n"
    "Path D (故障降级): 自检失败 → 进入 Degraded Mode → 强制位置灯 ON → DTC 上报\n"
    "Path E (手动超控): 驾驶员操作灯光开关 → 超控所有自动决策 → 直接驱动灯具",
    size=10.0,
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
#  4. BDD — Structure Pillar
# ═══════════════════════════════════════════════════════════════════════
add_heading_styled("4. BDD 块定义图 — 系统结构分解", level=1)
add_para("MagicGrid 定位: S3 · Solution Domain · Structure 列", bold=True, size=9.5)
add_para("BDD 将活动图中泳道代表的角色实例化为物理 Block，定义属性、类型和 Block 间的组合/引用关系。"
         "每个 Block 的 Value Property 直接来源于活动图的数据流和决策条件。", size=10.5)

add_para("从活动图到 BDD 的映射：", bold=True, size=10.5)
add_table(
    ["活动图泳道/元素", "→", "BDD Block", "关键属性"],
    [
        ["Swimlane: AED Controller", "→", "AedController", "ambientLightThreshold=1000 lux, debounceTime=2000ms, selfCheckTimeout=500ms"],
        ["Swimlane: Light Sensor", "→", "LightSensor", "sensitivity=1.0, filterWindow=50ms, outputRange=0..200000 lux"],
        ["Swimlane: Rain Sensor", "→", "RainSensor", "wipingLevel=0..4, sensitivity=0..5"],
        ["Swimlane: BCM", "→", "BodyControlModule (ref)", "signalUpdateRate=100Hz"],
        ["数据流: 环境光强", "→", "LightSensor.outputRange", "Real, 0..200000 lux"],
        ["执行: 点亮近光灯", "→", "Headlight", "nominalPower=55W, type=Halogen|LED|Xenon"],
        ["执行: 点亮日行灯", "→", "DaytimeRunningLight", "nominalPower=15W, dimmingEnabled=true"],
        ["决策: 手动超控", "→", "HeadlightSwitch (ref)", "position=OFF|AUTO|PARK|LOW|HIGH"],
    ],
)

add_image(INPUT_DIR / "AED_BDD.jpg", width_inches=9.0)
add_para("图 4-1   AED 系统块定义图 (Block Definition Diagram)", size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_para("结构设计原则：", bold=True, size=10.5)
add_para(
    "1. 组合关系 (Composition，黑色菱形): AedController *-- Sensor/Actuator — 传感器和灯具的生命周期由控制器管理。\n"
    "2. 引用关系 (Reference，空心菱形): AedController o-- Switch/BCM/Driver — 这些组件独立存在，AedController 只持有引用。\n"
    "3. 区分组合 vs 引用是 SysML v2 的核心概念——def 定义类型，usage 实例化。混淆两者会导致错误的资源管理策略。\n"
    "4. Headlight.feedbackCurrent 属性是实现故障检测的基础——BCM 周期性上报电流值，控制器与基准值比较判断断路/短路。",
    size=10.0,
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
#  5. 参数图 — Parameters Pillar
# ═══════════════════════════════════════════════════════════════════════
add_heading_styled("5. 参数图 — 约束方程与仲裁逻辑", level=1)
add_para("MagicGrid 定位: S4 · Solution Domain · Parameters 列", bold=True, size=9.5)
add_para("参数图将活动图中每个决策节点形式化为约束块（Constraint Block），"
         "并建立约束间的传播关系。BDD 中的 Value Property 作为约束的输入参数。"
         "整套约束网络可以直接转换为 Simulink/Modelica 模型执行仿真验证。", size=10.5)

add_para("从活动图决策节点到参数图约束块的映射：", bold=True, size=10.5)
add_table(
    ["活动图决策节点", "→", "参数图约束块", "约束方程式"],
    [
        ["光强 < 阈值_近光灯?", "→", "AmbientLightDecision", "ambientLight < 400 → LOW_BEAM; < 1000 → DRL; else → OFF"],
        ["雨量 > 阈值?", "→", "RainOverride", "forceLowBeam = (rainLevel >= 2)"],
        ["隧道检测?", "→", "SpeedOverride", "tunnelDetected = (speed > 0) AND (lightDropRate > 500 lux/s)"],
        ["手动超控?", "→", "ManualOverride", "autoModeActive = (switchPosition == AUTO)"],
        ["自检通过?", "→", "DegradationDecision", "degradedMode = !selfCheckPassed OR lightFaultActive"],
        ["综合仲裁", "→", "FinalArbitration", "Priority: Degraded > Manual > Rain > Auto"],
    ],
)

add_image(INPUT_DIR / "AED_Parametric.jpg", width_inches=9.0)
add_para("图 5-1   AED 灯光触发条件参数图 (Parametric Diagram)", size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_para("约束传播网络：", bold=True, size=10.5)
add_para(
    "  SpeedOverride ── influence ──→ AmbientLightDecision ── autoCommand ──┐\n"
    "  RainOverride ──── rainForce ──→                                       │\n"
    "  ManualOverride ─ autoModeActive ──→     FinalArbitration ←── degradedMode ── DegradationDecision\n"
    "                                                                        │\n"
    "                                                             LightingMode (枚举)\n\n"
    "FinalArbitration 是仲裁核心——它汇总 4 路输入，按优先级决定最终灯光模式。"
    "优先级设计反映了 ISO 26262 功能安全要求：降级模式 > 驾驶员意图 > 法规要求（雨天）> 自动决策。",
    size=10.0,
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
#  6. 时序图 — 跨视角交叉验证
# ═══════════════════════════════════════════════════════════════════════
add_heading_styled("6. 时序图 — 隧道场景全视角验证", level=1)
add_para("MagicGrid 定位: 交叉验证 · 实例化所有视角于一个具体运行场景", bold=True, size=9.5)
add_para("时序图选取「车辆驶入隧道→隧道内行驶→驶出隧道→近光灯断路故障」这一完整场景，"
         "将 BDD 的结构（参与者 = Block 实例）、活动图的流程（消息序列 = 活动执行顺序）、"
         "参数图的约束（条件判断 = 约束求值）全部实例化为一个可验证的交互序列。", size=10.5)

add_para("场景四阶段：", bold=True, size=10.5)
add_table(
    ["阶段", "场景", "关键交互", "验证的约束"],
    [
        ["1. 初始", "日间行驶, 80 km/h", "光照 50000 lux → DRL ON", "AmbientLightDecision: else → OFF; DRL 法规强制"],
        ["2. 入隧", "光照骤降 50000→350 lux", "Debounce 2000ms → 近光灯 ON, DRL Dim", "AmbientLightDecision: < 400 → LOW_BEAM; SpeedOverride: tunnelDetected"],
        ["3. 隧内", "稳定行驶, 光照 ~100 lux", "Heartbeat 无变化, 50ms 周期", "debounceTime=2000ms 防闪烁; 20Hz CAN 报文"],
        ["4. 出隧", "光照恢复 8000 lux", "Debounce 2000ms → 近光灯 OFF, DRL 恢复", "AmbientLightDecision: > 1000 → DRL"],
        ["5. 故障", "近光灯断路", "feedbackCurrent=0mA → DTC B1A32-11 → 位置灯 ON", "DegradationDecision: faultActive → safeState; ASIL-A 500ms 切换"],
    ],
)

add_image(INPUT_DIR / "AED_Sequence.jpg", width_inches=9.0)
add_para("图 6-1   AED 隧道场景时序图 (Sequence Diagram)", size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_para("时序图中的追溯验证：", bold=True, size=10.5)
add_para(
    "1. 每个 Lifeline（参与者）必须能在 BDD 中找到对应的 Block — 例如 AED Controller 对应 AedController，BCM 对应 BodyControlModule。\n"
    "2. 每个判断条件必须在参数图中找到对应的约束块 — 例如「光照 350lux < 400lux → 近光灯 ON」对应 AmbientLightDecision 约束。\n"
    "3. 每个故障响应必须在用例图中追溯到安全需求 — 例如断路→位置灯 ON 来自 UC10（故障诊断与降级）。\n"
    "4. debounce 2000ms 在 BDD（AedController.debounceTime）和时序图（两个 debounce 区间）中保持一致。",
    size=10.0,
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
#  7. 追溯矩阵
# ═══════════════════════════════════════════════════════════════════════
add_heading_styled("7. 完整追溯矩阵", level=1)
add_para("以下矩阵展示从需求（用例图）→ 行为（活动图）→ 结构（BDD）→ 参数（参数图）→ 交互（时序图）的完整追溯链：", size=10.5)

add_table(
    ["追溯链", "源 (Source)", "→", "目标 (Target)", "类型", "状态"],
    [
        ["R→B", "UC01 自动控制近光灯", "→", "Activity 主流程", "refine", "✓"],
        ["R→B", "UC08 手动超控", "→", "Activity 超控仲裁节点", "refine", "✓"],
        ["R→B", "UC10 故障诊断", "→", "Activity 故障检测分支", "refine", "✓"],
        ["B→S", "Swimlane: AED Controller", "→", "BDD: AedController block", "allocate", "✓"],
        ["B→S", "Swimlane: Light Sensor", "→", "BDD: LightSensor block", "allocate", "✓"],
        ["B→S", "数据流: 环境光强", "→", "BDD: LightSensor.outputRange", "allocate", "✓"],
        ["B→P", "决策: 光强 < 阈值?", "→", "PAR: AmbientLightDecision", "derive", "✓"],
        ["B→P", "决策: 雨量 > 阈值?", "→", "PAR: RainOverride", "derive", "✓"],
        ["B→P", "决策: 综合仲裁", "→", "PAR: FinalArbitration", "derive", "✓"],
        ["S→P", "BDD: ambientLightThreshold", "→", "PAR: AmbientLightDecision 输入", "bind", "✓"],
        ["S→P", "BDD: debounceTime", "→", "PAR: 时间约束条件", "bind", "✓"],
        ["B→I", "Activity 故障处理", "→", "SEQ: 断路场景阶段 5", "verify", "✓"],
        ["S→I", "BDD: Headlight.feedbackCurrent", "→", "SEQ: 电流监控消息", "verify", "✓"],
        ["P→I", "PAR: AmbientLightDecision", "→", "SEQ: 入隧/出隧判断", "verify", "✓"],
        ["P→I", "PAR: DegradationDecision", "→", "SEQ: 故障降级响应", "verify", "✓"],
    ],
)

add_para("追溯类型说明：", bold=True, size=10.5)
add_para(
    "  refine (精化): 用例→活动 — 把「做什么」展开为「怎么做」的操作序列\n"
    "  allocate (分配): 活动→结构 — 把行为职责分配到物理组件\n"
    "  derive (派生): 活动决策→参数约束 — 把经验判断形式化为可计算的方程式\n"
    "  bind (绑定): 结构属性→参数输入 — 把 BDD 的值属性绑定到约束变量\n"
    "  verify (验证): 时序图交叉检查 — 在一个具体场景中验证多视角一致性",
    size=9.5,
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════
#  8. 总结
# ═══════════════════════════════════════════════════════════════════════
add_heading_styled("8. 总结：MagicGrid 四支柱协同机制", level=1)

add_para(
    "本文档以 AED 自动外灯系统为例，展示了 MagicGrid 方法论的完整建模流程。"
    "五张图并非孤立存在，而是通过命名约定、元素映射和显式追溯关系编织成一张系统工程网：",
    size=10.5,
)

add_para(
    "1. 用例图（Requirements）锚定功能边界 — 定义了「做什么」的契约，是所有下游分析的起点。\n"
    "2. 活动图（Behavior）将核心用例展开为操作流程 — 决策节点成为参数图的约束来源，泳道角色成为 BDD 的 Block 来源。\n"
    "3. BDD（Structure）将行为角色实例化为物理组件 — 值属性（Value Property）直接绑定到参数图的约束输入变量。\n"
    "4. 参数图（Parameters）将经验规则形式化为约束方程 — 约束网络可直接导入 Simulink 仿真验证效能指标（MoE）。\n"
    "5. 时序图（Interaction）将以上全部视角实例化于一个具体场景 — 验证结构、行为、参数在运行时的时序一致性。",
    size=10.0,
)

add_para(
    "这种逐层精化、交叉验证的方法，保证了从利益相关者需求到软件实现的完整追溯性，"
    "避免「架构漂移」——即设计文档与实际实现之间日益扩大的鸿沟。"
    "在汽车电子领域，这也是满足 ISO 26262 功能安全标准和 ASPICE CL2/CL3 追溯性要求的基础实践。",
    size=10.0,
)

doc.add_paragraph()
add_para("— End —", size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ─────────────────────────────────────────────────────────────
OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT_PATH))
print(f"Document saved: {OUTPUT_PATH}")
print(f"File size: {OUTPUT_PATH.stat().st_size:,} bytes")
